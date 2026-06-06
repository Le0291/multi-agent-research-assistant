"""
file_processor.py — Extract clean text from uploaded files.

Supported formats:
  • .txt  — plain text, read directly
  • .pdf  — extract text page by page via pdfplumber
  • .docx — extract paragraphs via python-docx
  • .csv  — convert rows to readable text via pandas
  • .md   — plain text (Markdown)

Each extractor returns a plain string that can be fed directly into any
agent (NER, classification, analysis, writer, critic) as if it were a
scraped web page.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
#  Per-format extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_txt(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """Read plain text file; try utf-8 then latin-1 as fallback."""
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pdfplumber.
    Falls back to a simple byte-scan message if pdfplumber is unavailable.
    """
    try:
        import pdfplumber  # noqa: PLC0415
    except ImportError:
        logger.warning("pdfplumber not installed — PDF extraction unavailable.")
        return "[PDF content could not be extracted: install pdfplumber]"

    pages_text: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"--- Page {i} ---\n{text}")
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc)
        return f"[PDF extraction error: {exc}]"

    return "\n\n".join(pages_text) if pages_text else "[No extractable text found in PDF]"


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a Word .docx file using python-docx."""
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        logger.warning("python-docx not installed — DOCX extraction unavailable.")
        return "[DOCX content could not be extracted: install python-docx]"

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        return f"[DOCX extraction error: {exc}]"


def _extract_csv(file_bytes: bytes) -> str:
    """Convert CSV rows to a readable text block."""
    try:
        import pandas as pd  # noqa: PLC0415
        df = pd.read_csv(io.BytesIO(file_bytes))
        # Convert to a Markdown-style table for readability
        header  = " | ".join(str(c) for c in df.columns)
        divider = " | ".join(["---"] * len(df.columns))
        rows    = [" | ".join(str(v) for v in row) for _, row in df.iterrows()]
        return f"{header}\n{divider}\n" + "\n".join(rows[:200])  # cap at 200 rows
    except Exception as exc:
        logger.error("CSV extraction failed: %s", exc)
        return f"[CSV extraction error: {exc}]"


# ─────────────────────────────────────────────────────────────────────────────
#  Public interface
# ─────────────────────────────────────────────────────────────────────────────

# Map file extension → extractor function
_EXTRACTORS = {
    ".txt":  lambda b: _extract_txt(b),
    ".md":   lambda b: _extract_txt(b),
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".doc":  _extract_docx,
    ".csv":  _extract_csv,
}

SUPPORTED_EXTENSIONS = list(_EXTRACTORS.keys())
SUPPORTED_EXTENSIONS_DISPLAY = ", ".join(e.lstrip(".").upper() for e in SUPPORTED_EXTENSIONS)


def extract_text(file_bytes: bytes, filename: str, max_chars: int = 20_000) -> str:
    """
    Extract plain text from an uploaded file.

    Parameters
    ----------
    file_bytes : raw bytes from st.file_uploader
    filename   : original filename (used to detect format)
    max_chars  : truncate output to this length (avoids blowing context windows)

    Returns
    -------
    Extracted text string, or an error message if extraction failed.
    """
    ext = Path(filename).suffix.lower()
    extractor = _EXTRACTORS.get(ext)

    if not extractor:
        supported = SUPPORTED_EXTENSIONS_DISPLAY
        return (
            f"[Unsupported file type: '{ext}'.  "
            f"Supported: {supported}]"
        )

    logger.info("Extracting text from '%s' (format=%s, size=%d bytes)", filename, ext, len(file_bytes))
    text = extractor(file_bytes)

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[... truncated to {max_chars:,} characters ...]"

    return text


def get_file_info(file_bytes: bytes, filename: str) -> dict:
    """Return metadata about an uploaded file (for display in the UI)."""
    ext  = Path(filename).suffix.lower()
    size = len(file_bytes)

    size_str = (
        f"{size / 1024:.1f} KB"  if size < 1_000_000 else
        f"{size / 1_048_576:.1f} MB"
    )

    return {
        "name":      filename,
        "extension": ext.lstrip(".").upper(),
        "size":      size_str,
        "supported": ext in _EXTRACTORS,
    }
